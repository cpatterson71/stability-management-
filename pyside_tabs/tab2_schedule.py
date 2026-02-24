from PySide6.QtWidgets import (
    QWidget,
    QVBoxLayout,
    QFormLayout,
    QDateEdit,
    QGroupBox,
    QPushButton,
    QTableView,
    QLabel,
    QHeaderView,
    QHBoxLayout,
    QMessageBox,
    QScrollArea,
    QFrame,
    QFileDialog
)
from PySide6.QtCore import QDate, Qt, QAbstractTableModel
import pandas as pd
import json
import io
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

class PandasModel(QAbstractTableModel):
    """A model to interface a pandas DataFrame with QTableView."""
    def __init__(self, data):
        super().__init__()
        self._data = data

    def rowCount(self, parent=None):
        return self._data.shape[0]

    def columnCount(self, parent=None):
        return self._data.shape[1]

    def data(self, index, role=Qt.ItemDataRole.DisplayRole):
        if index.isValid():
            if role == Qt.ItemDataRole.DisplayRole:
                return str(self._data.iloc[index.row(), index.column()])
        return None

    def headerData(self, col, orientation, role):
        if orientation == Qt.Orientation.Horizontal and role == Qt.ItemDataRole.DisplayRole:
            return self._data.columns[col]
        return None

class Tab2Schedule(QWidget):
    def __init__(self, conn):
        super().__init__()
        self.conn = conn
        self.schedule_df = pd.DataFrame()
        self.setLayout(QVBoxLayout())

        # --- Filter Section ---
        filter_groupbox = QGroupBox("Filter Stability Testing Plan")
        filter_layout = QFormLayout()

        self.start_date_input = QDateEdit(QDate.currentDate().addYears(-1))
        self.start_date_input.setCalendarPopup(True)
        self.end_date_input = QDateEdit(QDate.currentDate().addYears(1))
        self.end_date_input.setCalendarPopup(True)
        
        filter_layout.addRow("Start Date:", self.start_date_input)
        filter_layout.addRow("End Date:", self.end_date_input)

        search_button = QPushButton("Search by Date Range")
        search_button.clicked.connect(self.search_schedule)
        filter_layout.addRow(search_button)

        filter_groupbox.setLayout(filter_layout)
        self.layout().addWidget(filter_groupbox)

        # --- Stability Pull Schedule ---
        schedule_groupbox = QGroupBox("Stability Pull Schedule")
        schedule_layout = QVBoxLayout()
        
        self.schedule_table = QTableView()
        schedule_layout.addWidget(self.schedule_table)

        schedule_groupbox.setLayout(schedule_layout)
        self.layout().addWidget(schedule_groupbox)

        # --- Stability Testing Plan ---
        plan_groupbox = QGroupBox("Stability Testing Plan")
        self.plan_layout = QVBoxLayout() # Made an attribute to access it later
        self.plan_layout.addWidget(QLabel("Testing plan will be displayed here.")) # Placeholder
        plan_groupbox.setLayout(self.plan_layout)
        self.layout().addWidget(plan_groupbox)

        # --- Document Generation ---
        doc_gen_groupbox = QGroupBox("Generate Request Documents")
        doc_gen_layout = QVBoxLayout()
        generate_button = QPushButton("Generate Documents")
        generate_button.clicked.connect(self.generate_documents)
        doc_gen_layout.addWidget(generate_button)
        doc_gen_groupbox.setLayout(doc_gen_layout)
        self.layout().addWidget(doc_gen_groupbox)

    def generate_documents(self):
        if self.schedule_df.empty:
            QMessageBox.warning(self, "No Data", "Please search for a schedule first.")
            return

        try:
            # Get master test info again, in case it's needed
            master_tests_df = pd.read_sql_query('SELECT test_name, test_method, form_no FROM master_tests', self.conn)
            test_procedure_map = master_tests_df.set_index('test_name').to_dict('index')

            grouped_by_client = self.schedule_df.groupby('Client Code')
            for client_code, client_df in grouped_by_client:
                doc = docx.Document()
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Arial'
                font.size = Pt(11)

                doc.add_heading(f"Stability Request for: {client_code}", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

                study_details = client_df[['Description', 'Protocol No.', 'Revision', 'Specification No.']].drop_duplicates().iloc[0]
                doc.add_paragraph(f"Description: {study_details['Description']}")
                doc.add_paragraph(f"Protocol: {study_details['Protocol No.']} Rev. {study_details['Revision']}")
                doc.add_paragraph(f"Specification: {study_details['Specification No.']}")
                doc.add_paragraph("Need by date:")
                doc.add_paragraph("Requestor Initials/Date:")
                doc.add_paragraph()

                doc.add_heading("Stability Pull Schedule", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
                docx_schedule_cols = ['Storage Condition', 'Lot No.', 'Timepoint', 'Pull Date', 'Number of Vials']
                schedule_table_df = client_df[docx_schedule_cols].drop_duplicates()
                
                table = doc.add_table(rows=1, cols=len(docx_schedule_cols))
                table.style = 'Table Grid'
                for i, col_name in enumerate(docx_schedule_cols):
                    table.cell(0, i).text = col_name
                for _, row in schedule_table_df.iterrows():
                    row_cells = table.add_row().cells
                    for i, col_name in enumerate(docx_schedule_cols):
                        row_cells[i].text = str(row[col_name])
                
                doc.add_paragraph()
                doc.add_heading("Consolidated Testing Plan", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                client_all_tests = set()
                for tests_str in client_df['tests_to_perform']:
                    if pd.notna(tests_str) and tests_str.strip():
                        client_all_tests.update(json.loads(tests_str))
                
                if client_all_tests:
                    test_plan_data = [{
                        "Test": test, "Test Method": test_procedure_map.get(test, {}).get('test_method', "N/A"),
                        "Form #": test_procedure_map.get(test, {}).get('form_no', "N/A"), "Copies": ""
                    } for test in sorted(list(client_all_tests))]
                    test_plan_df = pd.DataFrame(test_plan_data)
                    
                    test_table_cols = ["Test", "Test Method", "Form #", "Copies"]
                    test_table = doc.add_table(rows=1, cols=len(test_table_cols))
                    test_table.style = 'Table Grid'
                    for i, col_name in enumerate(test_table_cols):
                        test_table.cell(0, i).text = col_name
                    for _, row in test_plan_df.iterrows():
                        row_cells = test_table.add_row().cells
                        for i, col_name in enumerate(test_table_cols):
                            row_cells[i].text = str(row[col_name])
                
                # Prompt user to save the file
                start_date_str = self.start_date_input.date().toString("yyyy-MM-dd")
                end_date_str = self.end_date_input.date().toString("yyyy-MM-dd")
                default_filename = f"stability_request_{client_code}_{start_date_str}_to_{end_date_str}.docx"
                
                file_path, _ = QFileDialog.getSaveFileName(self, f"Save Document for {client_code}", default_filename, "Word Documents (*.docx)")
                
                if file_path:
                    doc.save(file_path)
                    QMessageBox.information(self, "Success", f"Document for {client_code} saved.")

            QMessageBox.information(self, "Complete", "All documents have been generated.")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate documents: {e}")


    def search_schedule(self):
        start_date = self.start_date_input.date().toString("yyyy-MM-dd")
        end_date = self.end_date_input.date().toString("yyyy-MM-dd")

        try:
            sql_query = """
                SELECT
                    tti.id,
                    ss.client_code AS "Client Code",
                    ss.description AS "Description",
                    ss.protocol_no AS "Protocol No.",
                    ss.revision AS "Revision",
                    ss.specification_no AS "Specification No.",
                    ss.lot_number AS "Lot No.",
                    sc.storage_condition AS "Storage Condition",
                    tti.timepoint AS "Timepoint",
                    tti.pull_date AS "Pull Date",
                    tti.num_vials AS "Number of Vials",
                    tti.num_copies AS "Number of Copies",
                    tti.tests_to_perform
                FROM timepoint_testing_info tti
                JOIN storage_schedules sc ON tti.schedule_id = sc.id
                JOIN stability_studies ss ON sc.study_id = ss.id
                WHERE tti.pull_date BETWEEN ? AND ?
                ORDER BY ss.client_code, ss.protocol_no, sc.storage_condition, tti.pull_date;
            """
            self.schedule_df = pd.read_sql_query(sql_query, self.conn, params=(start_date, end_date))

            if self.schedule_df.empty:
                QMessageBox.information(self, "No Results", "No stability pulls scheduled for the selected date range.")
                # Clear table view
                self.schedule_table.setModel(None)
            else:
                QMessageBox.information(self, "Success", f"Found {len(self.schedule_df)} stability pulls.")
                
                # Display schedule in the table
                display_cols = [
                    'Client Code', 'Description', 'Storage Condition', 'Protocol No.', 'Revision', 
                    'Specification No.', 'Lot No.', 'Timepoint', 'Pull Date', 'Number of Vials', 'Number of Copies'
                ]
                model = PandasModel(self.schedule_df[display_cols])
                self.schedule_table.setModel(model)
                self.schedule_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
                
                # Update the testing plan view
                self.update_testing_plan_display()

        except Exception as e:
            QMessageBox.critical(self, "Database Error", f"An error occurred during search: {e}")
            self.schedule_table.setModel(None) # Clear table on error

    def update_testing_plan_display(self):
        # Clear previous content
        for i in reversed(range(self.plan_layout.count())): 
            widget = self.plan_layout.itemAt(i).widget()
            if widget is not None:
                widget.deleteLater()

        if self.schedule_df.empty:
            self.plan_layout.addWidget(QLabel("Testing plan will be displayed here."))
            return
            
        try:
            # Get master test info for the test plan
            master_tests_df = pd.read_sql_query('SELECT test_name, test_method, form_no FROM master_tests', self.conn)
            test_procedure_map = master_tests_df.set_index('test_name').to_dict('index')
        except Exception as e:
            QMessageBox.warning(self, "DB Warning", f"Could not load master tests for test plan: {e}")
            test_procedure_map = {}
        
        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        scroll_content = QWidget()
        scroll_layout = QVBoxLayout(scroll_content)
        
        grouped_by_study = self.schedule_df.groupby(['Client Code', 'Protocol No.', 'Revision'])
        for (client_code, protocol_no, revision), study_df in grouped_by_study:
            scroll_layout.addWidget(QLabel(f"<b>Client Code:</b> {client_code}"))
            scroll_layout.addWidget(QLabel(f"<b>Protocol No.:</b> {protocol_no} (Rev. {revision})"))
            
            grouped_by_condition = study_df.groupby('Storage Condition')
            for condition, condition_df in grouped_by_condition:
                condition_label = QLabel(f"<b>Storage Condition:</b> {condition}")
                scroll_layout.addWidget(condition_label)
                
                all_tests = set()
                for tests_str in condition_df['tests_to_perform']:
                    if pd.notna(tests_str) and tests_str.strip():
                        try:
                            all_tests.update(json.loads(tests_str))
                        except json.JSONDecodeError:
                            pass # Ignore errors in this display logic

                if not all_tests:
                    scroll_layout.addWidget(QLabel("No tests scheduled for this condition."))
                else:
                    test_plan_data = [{
                        "Test": test, 
                        "Test Method": test_procedure_map.get(test, {}).get('test_method', "N/A"),
                        "Form #": test_procedure_map.get(test, {}).get('form_no', "N/A"),
                        "Copies": ""
                    } for test in sorted(list(all_tests))]
                    
                    plan_df = pd.DataFrame(test_plan_data)
                    plan_table = QTableView()
                    plan_model = PandasModel(plan_df)
                    plan_table.setModel(plan_model)
                    plan_table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Stretch)
                    plan_table.setFixedHeight(plan_df.shape[0] * 30 + 35) # Adjust height dynamically
                    scroll_layout.addWidget(plan_table)
            
            separator = QFrame()
            separator.setFrameShape(QFrame.Shape.HLine)
            separator.setFrameShadow(QFrame.Shadow.Sunken)
            scroll_layout.addWidget(separator)

        scroll_area.setWidget(scroll_content)
        self.plan_layout.addWidget(scroll_area)
