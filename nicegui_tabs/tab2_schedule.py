from nicegui import ui, app
import pandas as pd
from datetime import date
import logging
import json
import io
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tabs.utils import sanitize_sheet_name

def tab2_schedule_ui(conn):
    app.storage.user.schedule_df = pd.DataFrame()
    
    # Placeholder for the testing plan display (to be created dynamically)
    testing_plan_container = ui.column().classes('w-full')

    with ui.column().classes('w-full'):
        ui.label('Stability Schedule').classes('text-2xl font-bold')

        # --- Filter Section ---
        with ui.card().classes('w-full'):
            ui.card_section().classes('bg-blue-grey-10 text-white').header().add_slot('default').set_text('Filter Stability Testing Plan')
            
            ui.date('Start Date', value=str(date.today().replace(year=date.today().year - 1))).bind_value(app.storage.user, 'schedule_start_date')
            ui.date('End Date', value=str(date.today().replace(year=date.today().year + 1))).bind_value(app.storage.user, 'schedule_end_date')

            schedule_table = ui.table(
                columns=[
                    {'name': 'Client Code', 'label': 'Client Code', 'field': 'Client Code', 'align': 'left'},
                    {'name': 'Description', 'label': 'Description', 'field': 'Description', 'align': 'left'},
                    {'name': 'Storage Condition', 'label': 'Storage Condition', 'field': 'Storage Condition', 'align': 'left'},
                    {'name': 'Protocol No.', 'label': 'Protocol No.', 'field': 'Protocol No.', 'align': 'left'},
                    {'name': 'Revision', 'label': 'Revision', 'field': 'Revision', 'align': 'left'},
                    {'name': 'Specification No.', 'label': 'Specification No.', 'field': 'Specification No.', 'align': 'left'},
                    {'name': 'Lot No.', 'label': 'Lot No.', 'field': 'Lot No.', 'align': 'left'},
                    {'name': 'Timepoint', 'label': 'Timepoint', 'field': 'Timepoint', 'align': 'left'},
                    {'name': 'Pull Date', 'label': 'Pull Date', 'field': 'Pull Date', 'align': 'left'},
                    {'name': 'Number of Vials', 'label': 'Number of Vials', 'field': 'Number of Vials', 'align': 'left'},
                    {'name': 'Number of Copies', 'label': 'Number of Copies', 'field': 'Number of Copies', 'align': 'left'},
                ],
                rows=[],
                row_key='id'
            ).classes('w-full')

            async def search_schedule():
                start_date_str = app.storage.user.schedule_start_date
                end_date_str = app.storage.user.schedule_end_date

                try:
                    sql_query = """
                        SELECT
                            tti.id,
                            ss.client_code AS "Client Code",
                            ss.description AS "Description",
                            ss.protocol_no AS "Protocol No.",
                            ss.revision AS "Revision",
                            ss.specification_no AS "Specification No.",
                            ss.lot_number AS "Lot No.",
                            sc.storage_condition AS "Storage Condition",
                            tti.timepoint AS "Timepoint",
                            tti.pull_date AS "Pull Date",
                            tti.num_vials AS "Number of Vials",
                            tti.num_copies AS "Number of Copies",
                            tti.tests_to_perform
                        FROM timepoint_testing_info tti
                        JOIN storage_schedules sc ON tti.schedule_id = sc.id
                        JOIN stability_studies ss ON sc.study_id = ss.id
                        WHERE tti.pull_date BETWEEN ? AND ?
                        ORDER BY ss.client_code, ss.protocol_no, sc.storage_condition, tti.pull_date;
                    """
                    app.storage.user.schedule_df = pd.read_sql_query(sql_query, conn, params=(start_date_str, end_date_str))

                    if app.storage.user.schedule_df.empty:
                        ui.notify("No stability pulls scheduled for the selected date range.", type='info')
                        schedule_table.rows = []
                    else:
                        ui.notify(f"Found {len(app.storage.user.schedule_df)} stability pulls.", type='positive')
                        schedule_table.rows = app.storage.user.schedule_df.to_dict('records')
                    schedule_table.update()

                    # Update the testing plan display
                    await update_testing_plan_display(testing_plan_container)

                except Exception as e:
                    ui.notify(f"An error occurred during search: {e}", type='negative')
                    schedule_table.rows = []
                    schedule_table.update()
            
            ui.button('Search by Date Range', on_click=search_schedule)

        # --- Stability Pull Schedule ---
        with ui.card().classes('w-full'):
            ui.card_section().classes('bg-blue-grey-10 text-white').header().add_slot('default').set_text('Stability Pull Schedule')
            ui.add_at(ui.get_slot_stack()[-1], schedule_table) # Add the already defined table here

        # --- Stability Testing Plan ---
        with ui.card().classes('w-full'):
            ui.card_section().classes('bg-blue-grey-10 text-white').header().add_slot('default').set_text('Stability Testing Plan')
            ui.add_at(ui.get_slot_stack()[-1], testing_plan_container) # Add the testing plan container here

        # --- Document Generation ---
        with ui.card().classes('w-full'):
            ui.card_section().classes('bg-blue-grey-10 text-white').header().add_slot('default').set_text('Generate Request Documents')
            
            @ui.button('Generate Documents').on('click')
            async def generate_documents():
                if app.storage.user.schedule_df.empty:
                    ui.notify("Please search for a schedule first.", type='warning')
                    return

                try:
                    master_tests_df = pd.read_sql_query('SELECT test_name, test_method, form_no FROM master_tests', conn)
                    test_procedure_map = master_tests_df.set_index('test_name').to_dict('index')

                    grouped_by_client = app.storage.user.schedule_df.groupby('Client Code')
                    for client_code, client_df in grouped_by_client:
                        doc = docx.Document()
                        style = doc.styles['Normal']
                        font = style.font
                        font.name = 'Arial'
                        font.size = Pt(11)

                        doc.add_heading(f"Stability Request for: {client_code}", level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER

                        study_details = client_df[['Description', 'Protocol No.', 'Revision', 'Specification No.']].drop_duplicates().iloc[0]
                        doc.add_paragraph(f"Description: {study_details['Description']}")
                        doc.add_paragraph(f"Protocol: {study_details['Protocol No.']} Rev. {study_details['Revision']}")
                        doc.add_paragraph(f"Specification: {study_details['Specification No.']}")
                        doc.add_paragraph("Need by date:")
                        doc.add_paragraph("Requestor Initials/Date:")
                        doc.add_paragraph()

                        doc.add_heading("Stability Pull Schedule", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
                        docx_schedule_cols = ['Storage Condition', 'Lot No.', 'Timepoint', 'Pull Date', 'Number of Vials']
                        schedule_table_df = client_df[docx_schedule_cols].drop_duplicates()
                        
                        table = doc.add_table(rows=1, cols=len(docx_schedule_cols))
                        table.style = 'Table Grid'
                        for i, col_name in enumerate(docx_schedule_cols):
                            table.cell(0, i).text = col_name
                        for _, row in schedule_table_df.iterrows():
                            row_cells = table.add_row().cells
                            for i, col_name in enumerate(docx_schedule_cols):
                                row_cells[i].text = str(row[col_name])
                        
                        doc.add_paragraph()
                        doc.add_heading("Consolidated Testing Plan", level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        client_all_tests = set()
                        for tests_str in client_df['tests_to_perform']:
                            if pd.notna(tests_str) and tests_str.strip():
                                client_all_tests.update(json.loads(tests_str))
                        
                        if client_all_tests:
                            test_plan_data = [{
                                "Test": test, "Test Method": test_procedure_map.get(test, {}).get('test_method', "N/A"),
                                "Form #": test_procedure_map.get(test, {}).get('form_no', "N/A"), "Copies": ""
                            } for test in sorted(list(client_all_tests))]
                            test_plan_df = pd.DataFrame(test_plan_data)
                            
                            test_table_cols = ["Test", "Test Method", "Form #", "Copies"]
                            test_table = doc.add_table(rows=1, cols=len(test_table_cols))
                            test_table.style = 'Table Grid'
                            for i, col_name in enumerate(test_table_cols):
                                test_table.cell(0, i).text = col_name
                            for _, row in test_plan_df.iterrows():
                                row_cells = test_table.add_row().cells
                                for i, col_name in enumerate(test_table_cols):
                                    row_cells[i].text = str(row[col_name])
                        
                        doc_io = io.BytesIO()
                        doc.save(doc_io)
                        doc_io.seek(0)
                        
                        start_date_str = app.storage.user.schedule_start_date
                        end_date_str = app.storage.user.schedule_end_date
                        file_name = f"stability_request_{client_code}_{start_date_str}_to_{end_date_str}.docx"
                        ui.download(doc_io.getvalue(), file_name)

                    ui.notify("Documents generated successfully!", type='positive')

                except Exception as e:
                    ui.notify(f"Failed to generate documents: {e}", type='negative')

    async def update_testing_plan_display(container):
        # Clear previous content
        container.clear()

        if app.storage.user.schedule_df.empty:
            with container:
                ui.label("Testing plan will be displayed here.")
            return

        try:
            master_tests_df = pd.read_sql_query('SELECT test_name, test_method, form_no FROM master_tests', conn)
            test_procedure_map = master_tests_df.set_index('test_name').to_dict('index')
        except Exception as e:
            ui.notify(f"Could not load master tests for test plan: {e}", type='warning')
            test_procedure_map = {}

        with container:
            grouped_by_study = app.storage.user.schedule_df.groupby(['Client Code', 'Protocol No.', 'Revision'])
            for (client_code, protocol_no, revision), study_df in grouped_by_study:
                ui.label(f"<b>Client Code:</b> {client_code}").props('html')
                ui.label(f"<b>Protocol No.:</b> {protocol_no} (Rev. {revision})").props('html')
                
                grouped_by_condition = study_df.groupby('Storage Condition')
                for condition, condition_df in grouped_by_condition:
                    ui.label(f"<b>Storage Condition:</b> {condition}").props('html')
                    
                    all_tests = set()
                    for tests_str in condition_df['tests_to_perform']:
                        if pd.notna(tests_str) and tests_str.strip():
                            try:
                                all_tests.update(json.loads(tests_str))
                            except json.JSONDecodeError:
                                pass # Ignore errors in this display logic

                    if not all_tests:
                        ui.label("No tests scheduled for this condition.")
                    else:
                        test_plan_data = [{
                            "Test": test, 
                            "Test Method": test_procedure_map.get(test, {}).get('test_method', "N/A"),
                            "Form #": test_procedure_map.get(test, {}).get('form_no', "N/A"),
                            "Copies": ""
                        } for test in sorted(list(all_tests))]
                        
                        test_plan_table = ui.table(
                            columns=[
                                {'name': 'Test', 'label': 'Test', 'field': 'Test', 'align': 'left'},
                                {'name': 'Test Method', 'label': 'Test Method', 'field': 'Test Method', 'align': 'left'},
                                {'name': 'Form #', 'label': 'Form #', 'field': 'Form #', 'align': 'left'},
                                {'name': 'Copies', 'label': 'Copies', 'field': 'Copies', 'align': 'left'},
                            ],
                            rows=test_plan_data,
                            row_key='Test',
                            pagination={'rowsPerPage': 0}
                        ).classes('w-full')
                ui.separator()
