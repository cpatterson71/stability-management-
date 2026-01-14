import streamlit as st
import pandas as pd
from datetime import datetime
from dateutil.relativedelta import relativedelta
import json
import logging
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def show_schedule_tab(conn):
    logging.info("Entering Tab 2: Stability Schedule")
    st.header("Stability Schedule")

    # Ensure all required session state variables are initialized
    if 'schedule_df' not in st.session_state:
        st.session_state.schedule_df = pd.DataFrame()
    if 'df_display' not in st.session_state:
        st.session_state.df_display = pd.DataFrame()
    if 'filtered_schedule_df' not in st.session_state:
        st.session_state.filtered_schedule_df = pd.DataFrame()
    if 'docs_to_download' not in st.session_state:
        st.session_state.docs_to_download = {}

    if conn:
        with st.expander("Query Schedule by Date Range", expanded=True):
            start_date = st.date_input("Start Date", value=datetime.today() - relativedelta(years=1))
            end_date = st.date_input("End Date", value=datetime.today() + relativedelta(years=1))
            
            if st.button("Query Schedule"):
                if start_date and end_date:
                    start_date_str = start_date.strftime('%Y-%m-%d')
                    end_date_str = end_date.strftime('%Y-%m-%d')
                    
                    sql_query = """
                        SELECT
                            tti.id,
                            ss.client_code AS 'Client Code',
                            ss.description AS 'Description',
                            ss.protocol_no AS 'Protocol No.',
                            ss.revision AS 'Revision',
                            ss.specification_no AS 'Specification No.',
                            ss.lot_number AS 'Lot No.',
                            sc.storage_condition AS 'Storage Condition',
                            tti.timepoint AS 'Timepoint',
                            tti.pull_date AS 'Pull Date',
                            tti.num_vials AS 'Number of Vials',
                            tti.num_copies AS 'Number of Copies',
                            tti.tests_to_perform
                        FROM timepoint_testing_info tti
                        JOIN storage_schedules sc ON tti.schedule_id = sc.id
                        JOIN stability_studies ss ON sc.study_id = ss.id
                        WHERE tti.pull_date BETWEEN %s AND %s
                        ORDER BY ss.client_code, ss.protocol_no, sc.storage_condition, tti.pull_date;
                    """
                    
                    # 1. Consistently populate schedule_df from the database
                    st.session_state.schedule_df = pd.read_sql_query(sql_query, conn, params=(start_date_str, end_date_str))
                    
                    # Reset dependent dataframes
                    st.session_state.df_display = pd.DataFrame()
                    st.session_state.filtered_schedule_df = pd.DataFrame()

                    if st.session_state.schedule_df.empty:
                        st.info("No stability pulls scheduled for the selected date range.")
                    else:
                        st.success(f"Found {len(st.session_state.schedule_df)} stability pulls for {len(st.session_state.schedule_df['Lot No.'].unique())} lot(s).")
                        
                        # 2. Always derive df_display directly from schedule_df
                        df_display = st.session_state.schedule_df.copy()
                        df_display['Tests to Perform'] = df_display['tests_to_perform'].apply(
                            lambda x: ", ".join(json.loads(x)) if pd.notna(x) and x.strip() else ""
                        )
                        st.session_state.df_display = df_display

        # This check is now robust because df_display is always initialized
        if not st.session_state.df_display.empty:
            st.subheader("Stability Pull Schedule")
            
            schedule_cols = [
                'Client Code', 'Description', 'Storage Condition', 'Protocol No.', 'Revision', 'Specification No.', 
                'Lot No.', 'Timepoint', 'Pull Date', 'Number of Vials', 'Number of Copies'
            ]
            
            st.dataframe(
                st.session_state.df_display[schedule_cols],
                key="schedule_data_viewer",
                hide_index=True,
            )

            st.markdown("---")

            st.subheader("Filter Stability Testing Plan")
            search_client_code = st.text_input("Client Code", key="search_client_code")
            search_description = st.text_input("Description", key="search_description")
            search_protocol_no = st.text_input("Protocol Number", key="search_protocol_no")
            search_lot_no = st.text_input("Lot Number", key="search_lot_no")

            if st.button("Search", key="filter_test_plan"):
                # 3. Filtering logic acts on df_display and stores in filtered_schedule_df
                filtered_df = st.session_state.df_display.copy()
                if search_client_code:
                    filtered_df = filtered_df[filtered_df['Client Code'].str.contains(search_client_code, case=False, na=False)]
                if search_description:
                    filtered_df = filtered_df[filtered_df['Description'].str.contains(search_description, case=False, na=False)]
                if search_protocol_no:
                    filtered_df = filtered_df[filtered_df['Protocol No.'].str.contains(search_protocol_no, case=False, na=False)]
                if search_lot_no:
                    filtered_df = filtered_df[filtered_df['Lot No.'].str.contains(search_lot_no, case=False, na=False)]
                st.session_state.filtered_schedule_df = filtered_df
            
            st.subheader("Stability Testing Plan")
            
            # 4. Display components and generation logic strictly refer to session state
            display_and_gen_df = st.session_state.filtered_schedule_df if not st.session_state.filtered_schedule_df.empty else st.session_state.df_display

            if not display_and_gen_df.empty:
                test_procedure_map = {}
                if 'master_tests_df' in st.session_state and not st.session_state.master_tests_df.empty:
                    test_procedure_map = st.session_state.master_tests_df.set_index('Test')[['Test Method', 'Form No']].to_dict('index')

                grouped_by_study = display_and_gen_df.groupby(['Client Code', 'Protocol No.', 'Revision'])
                for (client_code, protocol_no, revision), study_df in grouped_by_study:
                    st.markdown(f"**Client Code:** `{client_code}`")
                    st.markdown(f"**Protocol No.:** `{protocol_no}` (Rev. `{revision}`)")
                    grouped_by_condition = study_df.groupby('Storage Condition')
                    for condition, condition_df in grouped_by_condition:
                        with st.container():
                            st.markdown(f"**Storage Condition:** {condition}")
                            
                            all_tests = set()
                            # Use the original 'tests_to_perform' from schedule_df for accuracy
                            original_indices = condition_df.index
                            original_tests_series = st.session_state.schedule_df.loc[original_indices, 'tests_to_perform']
                            
                            for tests_str in original_tests_series:
                                if pd.notna(tests_str) and tests_str.strip():
                                    try:
                                        tests = json.loads(tests_str)
                                        all_tests.update(tests)
                                    except json.JSONDecodeError:
                                        # Handle cases where the string might not be valid JSON
                                        logging.warning(f"Could not decode JSON for tests: {tests_str}")

                            if not all_tests:
                                st.text("No tests scheduled for this condition in the selected date range.")
                            else:
                                test_plan_data = [{
                                    "Test": test, 
                                    "Test Method": test_procedure_map.get(test, {}).get('Test Method', "N/A"),
                                    "Form #": test_procedure_map.get(test, {}).get('Form No', "N/A"),
                                    "Copies": "" # Blank as requested
                                } for test in sorted(list(all_tests))]
                                st.dataframe(pd.DataFrame(test_plan_data), hide_index=True, width='stretch')
                    st.markdown("---")

            # --- Document Generation Section ---
            st.subheader("Generate Request Documents")
            if st.button("Generate Documents"):
                try:
                    st.session_state.docs_to_download = {}
                    # Use the same logic for which df to use for generation
                    doc_gen_df = st.session_state.filtered_schedule_df if not st.session_state.filtered_schedule_df.empty else st.session_state.df_display
                    
                    if not doc_gen_df.empty:
                        grouped_by_client = doc_gen_df.groupby('Client Code')
                        for client_code, client_df in grouped_by_client:
                            doc = docx.Document()
                            
                            style = doc.styles['Normal']
                            font = style.font
                            font.name = 'Arial'
                            font.size = Pt(11)

                            heading = doc.add_heading(f"Stability Request for: {client_code}", level=1)
                            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

                            study_details = client_df[['Description', 'Protocol No.', 'Revision', 'Specification No.']].drop_duplicates().iloc[0]
                            doc.add_paragraph(f"Description: {study_details['Description']}")
                            doc.add_paragraph(f"Protocol: {study_details['Protocol No.']} Rev. {study_details['Revision']}")
                            doc.add_paragraph(f"Specification: {study_details['Specification No.']}")
                            doc.add_paragraph("Need by date:")
                            doc.add_paragraph()
                            doc.add_paragraph("Requestor Initials/Date:")

                            heading = doc.add_heading("Stability Pull Schedule", level=2)
                            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            docx_schedule_cols = ['Storage Condition', 'Lot No.', 'Timepoint', 'Pull Date', 'Number of Vials']
                            schedule_table_df = client_df[docx_schedule_cols].drop_duplicates()
                            
                            table = doc.add_table(rows=1, cols=len(docx_schedule_cols))
                            table.style = 'Table Grid'
                            hdr_cells = table.rows[0].cells
                            for i, col_name in enumerate(docx_schedule_cols):
                                hdr_cells[i].text = col_name
                                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for index, row in schedule_table_df.iterrows():
                                row_cells = table.add_row().cells
                                for i, col_name in enumerate(docx_schedule_cols):
                                    row_cells[i].text = str(row[col_name])
                                    row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                            doc.add_paragraph()
                            heading = doc.add_heading("Consolidated Testing Plan", level=2)
                            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            # Get all tests from the original, unfiltered data for this client
                            original_indices = client_df.index
                            original_tests_series = st.session_state.schedule_df.loc[original_indices, 'tests_to_perform']
                            client_all_tests = set()
                            for tests_str in original_tests_series:
                                if pd.notna(tests_str) and tests_str.strip():
                                    client_all_tests.update(json.loads(tests_str))
                            
                            if client_all_tests:
                                test_plan_list = []
                                for test in sorted(list(client_all_tests)):
                                    test_info = test_procedure_map.get(test, {})
                                    test_plan_list.append({
                                        "Test": test,
                                        "Test Method": test_info.get('Test Method', "N/A"),
                                        "Form #": test_info.get('Form No', "N/A"),
                                        "Copies": ""
                                    })
                                test_plan_df = pd.DataFrame(test_plan_list)

                                test_table_cols = ["Test", "Test Method", "Form #", "Copies"]
                                test_table = doc.add_table(rows=1, cols=len(test_table_cols))
                                test_table.style = 'Table Grid'
                                hdr_cells = test_table.rows[0].cells
                                for i, col_name in enumerate(test_table_cols):
                                    hdr_cells[i].text = col_name
                                    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                                for index, row in test_plan_df.iterrows():
                                    row_cells = test_table.add_row().cells
                                    for i, col_name in enumerate(test_table_cols):
                                        row_cells[i].text = str(row[col_name])
                                        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            else:
                                doc.add_paragraph("No tests scheduled for this client in the selected date range.")

                            doc_io = io.BytesIO()
                            doc.save(doc_io)
                            doc_io.seek(0)
                            st.session_state.docs_to_download[client_code] = doc_io.getvalue()
                        
                        if not st.session_state.docs_to_download:
                            st.warning("No documents were generated.")
                        st.rerun()
                    else:
                        st.warning("No data to generate documents from. Please query a schedule first.")

                except Exception as e:
                    st.error(f"Failed to generate document: {e}")
                    logging.error(f"Error during DOCX generation: {e}", exc_info=True)
            
            if st.session_state.get('docs_to_download'):
                st.markdown("---")
                st.info("Your documents are ready to download:")
                for client_code, doc_bytes in st.session_state.docs_to_download.items():
                    st.download_button(
                        label=f"📥 Download DOCX for {client_code}",
                        data=doc_bytes,
                        file_name=f"stability_request_{client_code}_{start_date.strftime('%Y-%m-%d')}_to_{end_date.strftime('%Y-%m-%d')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_docx_{client_code}"
                    )
                if st.button("Clear Generated Documents"):
                    st.session_state.docs_to_download = {}
                    st.rerun()
    else:
        st.error("Database connection is not available.")
        logging.error("Tab 2: Database connection not available.")